import os
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


class DocumentService:
    """Service class for Word document operations."""

    async def create_document(self, filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
        """Create a new Word document with optional metadata.
        
        Args:
            filename: Name of the document to create (with or without .docx extension)
            title: Optional title for the document metadata
            author: Optional author for the document metadata
        """
        filename = self._ensure_docx_extension(filename)
            
        try:
            doc = Document()
            
            # Set properties if provided
            if title:
                doc.core_properties.title = title
            if author:
                doc.core_properties.author = author
            
            # Ensure necessary styles exist
            self._ensure_heading_style(doc)
            self._ensure_table_style(doc)
            
            # Save the document
            doc.save(filename)
            
            return f"Document {filename} created successfully"
        except Exception as e:
            return f"Failed to create document: {str(e)}"

    async def add_heading(self, filename: str, text: str, level: int = 1) -> str:
        """Add a heading to a Word document.
        
        Args:
            filename: Path to the Word document
            text: Heading text
            level: Heading level (1-9, where 1 is the highest level)
        """
        filename = self._ensure_docx_extension(filename)
        
        # Ensure level is converted to integer
        try:
            level = int(level)
        except (ValueError, TypeError):
            return "Invalid parameter: level must be an integer between 1 and 9"
        
        # Validate level range
        if level < 1 or level > 9:
            return f"Invalid heading level: {level}. Level must be between 1 and 9."
        
        if not os.path.exists(filename):
            return f"Document {filename} does not exist"
        
        try:
            doc = Document(filename)
            
            # Ensure heading styles exist
            self._ensure_heading_style(doc)
            
            # Try to add heading with style
            try:
                heading = doc.add_heading(text, level=level)
                doc.save(filename)
                return f"Heading '{text}' (level {level}) added to {filename}"
            except Exception as style_error:
                # If style-based approach fails, use direct formatting
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles['Normal']
                run = paragraph.runs[0]
                run.bold = True
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
                
                doc.save(filename)
                return f"Heading '{text}' added to {filename} with direct formatting (style not available)"
        except Exception as e:
            return f"Failed to add heading: {str(e)}"

    async def add_paragraph(self, filename: str, text: str, style: Optional[str] = None) -> str:
        """Add a paragraph to a Word document.
        
        Args:
            filename: Path to the Word document
            text: Paragraph text
            style: Optional paragraph style name
        """
        filename = self._ensure_docx_extension(filename)
        
        if not os.path.exists(filename):
            return f"Document {filename} does not exist"
            
        try:
            doc = Document(filename)
            paragraph = doc.add_paragraph(text)
            
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    # Style doesn't exist, use normal and report it
                    paragraph.style = doc.styles['Normal']
                    doc.save(filename)
                    return f"Style '{style}' not found, paragraph added with default style to {filename}"
            
            doc.save(filename)
            return f"Paragraph added to {filename}"
        except Exception as e:
            return f"Failed to add paragraph: {str(e)}"


    @staticmethod
    def _ensure_docx_extension(filename: str) -> str:
        """
        Ensure filename has .docx extension.
        
        Args:
            filename: The filename to check
            
        Returns:
            Filename with .docx extension
        """
        if not filename.endswith('.docx'):
            return filename + '.docx'
        return filename

    @staticmethod
    def _ensure_heading_style(doc):
        """
        Ensure Heading styles exist in the document.
        
        Args:
            doc: Document object
        """
        for i in range(1, 10):  # Create Heading 1 through Heading 9
            style_name = f'Heading {i}'
            try:
                # Try to access the style to see if it exists
                style = doc.styles[style_name]
            except KeyError:
                # Create the style if it doesn't exist
                try:
                    style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    if i == 1:
                        style.font.size = Pt(16)
                        style.font.bold = True
                    elif i == 2:
                        style.font.size = Pt(14)
                        style.font.bold = True
                    else:
                        style.font.size = Pt(12)
                        style.font.bold = True
                except Exception:
                    # If style creation fails, we'll just use default formatting
                    pass

    @staticmethod
    def _ensure_table_style(doc):
        """
        Ensure Table Grid style exists in the document.
        
        Args:
            doc: Document object
        """
        try:
            # Try to access the style to see if it exists
            style = doc.styles['Table Grid']
        except KeyError:
            # If style doesn't exist, we'll handle it at usage time
            pass
